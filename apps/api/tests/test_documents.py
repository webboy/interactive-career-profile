import io
from unittest.mock import patch

import pytest
from docx import Document as DocxDocument
from httpx import AsyncClient
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.enums import DocumentFileType, DocumentSourceType, DocumentStatus, EmbeddingStatus, Visibility
from app.db.models.document import Document, DocumentChunk
from app.services.chunking import chunk_text
from app.services.documents import list_public_document_chunks
from app.services.extraction import extract_text_from_bytes


TEXT_PAYLOAD = {
    "title": "Career Notes",
    "content": "Led platform engineering initiatives across multiple teams.",
}

def _make_docx_bytes(text: str) -> bytes:
    document = DocxDocument()
    document.add_paragraph(text)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


@pytest.mark.asyncio
async def test_documents_require_auth(client: AsyncClient) -> None:
    response = await client.get("/api/admin/documents")
    assert response.status_code == 401


def test_chunk_text_overlap() -> None:
    text = "abcdefghijklmnopqrstuvwxyz" * 10
    segments = chunk_text(text, chunk_size_chars=50, chunk_overlap_chars=10)
    assert len(segments) > 1
    assert segments[0].char_start == 0
    assert segments[1].char_start == segments[0].char_end - 10


def test_extract_docx_text() -> None:
    content = _make_docx_bytes("Platform engineering experience")
    extracted = extract_text_from_bytes(DocumentFileType.DOCX, content)
    assert "Platform engineering experience" in extracted


@pytest.mark.asyncio
async def test_text_document_create_extract_and_chunk(auth_client: AsyncClient) -> None:
    create_response = await auth_client.post("/api/admin/documents/text", json=TEXT_PAYLOAD)
    assert create_response.status_code == 201
    created = create_response.json()
    assert created["source_type"] == DocumentSourceType.PASTED_TEXT
    assert created["visibility"] == Visibility.DRAFT
    assert created["status"] == DocumentStatus.EXTRACTED
    document_id = created["id"]

    chunk_response = await auth_client.post(f"/api/admin/documents/{document_id}/chunk")
    assert chunk_response.status_code == 200
    chunked = chunk_response.json()["document"]
    assert chunked["status"] == DocumentStatus.CHUNKED
    assert chunk_response.json()["chunks_created"] >= 1

    chunks_response = await auth_client.get(f"/api/admin/documents/{document_id}/chunks")
    assert chunks_response.status_code == 200
    chunks = chunks_response.json()
    assert len(chunks) >= 1
    assert chunks[0]["visibility"] == Visibility.DRAFT


@pytest.mark.asyncio
async def test_upload_txt_extract_and_chunk(auth_client: AsyncClient, tmp_path) -> None:
    with patch("app.services.documents.get_storage_driver") as mock_driver_factory:
        from app.services.storage.local import LocalStorageDriver

        mock_driver_factory.return_value = LocalStorageDriver(str(tmp_path))

        upload_response = await auth_client.post(
            "/api/admin/documents/upload",
            files={"file": ("resume.txt", b"Built RAG systems and APIs.", "text/plain")},
            data={"title": "Resume"},
        )
        assert upload_response.status_code == 201
        uploaded = upload_response.json()
        assert uploaded["original_filename"] == "resume.txt"
        assert uploaded["status"] == DocumentStatus.UPLOADED
        document_id = uploaded["id"]

        extract_response = await auth_client.post(f"/api/admin/documents/{document_id}/extract")
        assert extract_response.status_code == 200
        extracted = extract_response.json()["document"]
        assert extracted["status"] == DocumentStatus.EXTRACTED
        assert "RAG systems" in extracted["extracted_text"]

        chunk_response = await auth_client.post(f"/api/admin/documents/{document_id}/chunk")
        assert chunk_response.status_code == 200
        assert chunk_response.json()["chunks_created"] >= 1


@pytest.mark.asyncio
async def test_upload_rejects_oversized_file(auth_client: AsyncClient, tmp_path) -> None:
    class SmallLimitSettings:
        document_upload_max_bytes = 10
        local_storage_path = str(tmp_path)
        filesystem_driver = "local"

    with patch("app.services.documents.get_settings", return_value=SmallLimitSettings()):
        with patch("app.services.documents.get_storage_driver") as mock_driver_factory:
            from app.services.storage.local import LocalStorageDriver

            mock_driver_factory.return_value = LocalStorageDriver(str(tmp_path))

            response = await auth_client.post(
                "/api/admin/documents/upload",
                files={"file": ("large.txt", b"x" * 20, "text/plain")},
            )
            assert response.status_code == 413


@pytest.mark.asyncio
async def test_retry_ingestion_and_embedding_placeholder(auth_client: AsyncClient) -> None:
    create_response = await auth_client.post("/api/admin/documents/text", json=TEXT_PAYLOAD)
    document_id = create_response.json()["id"]

    embed_before_chunk = await auth_client.post(
        f"/api/admin/documents/{document_id}/request-embedding"
    )
    assert embed_before_chunk.status_code == 200
    assert embed_before_chunk.json()["embedding_status"] == EmbeddingStatus.FAILED

    retry_response = await auth_client.post(f"/api/admin/documents/{document_id}/retry-ingestion")
    assert retry_response.status_code == 200
    retried = retry_response.json()["document"]
    assert retried["status"] == DocumentStatus.CHUNKED
    assert retry_response.json()["chunks_created"] >= 1

    embed_after_chunk = await auth_client.post(
        f"/api/admin/documents/{document_id}/request-embedding"
    )
    assert embed_after_chunk.status_code == 200
    assert embed_after_chunk.json()["embedding_status"] == EmbeddingStatus.PENDING


@pytest.mark.asyncio
async def test_update_document_and_chunk_visibility(auth_client: AsyncClient) -> None:
    create_response = await auth_client.post("/api/admin/documents/text", json=TEXT_PAYLOAD)
    document_id = create_response.json()["id"]

    update_response = await auth_client.put(
        f"/api/admin/documents/{document_id}",
        json={"title": "Updated Notes", "visibility": Visibility.PUBLIC},
    )
    assert update_response.status_code == 200
    assert update_response.json()["title"] == "Updated Notes"
    assert update_response.json()["visibility"] == Visibility.PUBLIC

    await auth_client.post(f"/api/admin/documents/{document_id}/chunk")
    chunks_response = await auth_client.get(f"/api/admin/documents/{document_id}/chunks")
    chunk_id = chunks_response.json()[0]["id"]

    chunk_update = await auth_client.put(
        f"/api/admin/document-chunks/{chunk_id}",
        json={"visibility": Visibility.PUBLIC},
    )
    assert chunk_update.status_code == 200
    assert chunk_update.json()["visibility"] == Visibility.PUBLIC


@pytest.mark.asyncio
async def test_public_document_chunk_filter(db_session: AsyncSession) -> None:
    document = Document(
        title="Public Doc",
        source_type=DocumentSourceType.PASTED_TEXT,
        file_type=DocumentFileType.TEXT,
        extracted_text="Public and private chunks",
        visibility=Visibility.DRAFT,
        status=DocumentStatus.EXTRACTED,
    )
    db_session.add(document)
    await db_session.flush()

    db_session.add_all(
        [
            DocumentChunk(
                document_id=document.id,
                chunk_index=0,
                content="Public chunk",
                char_start=0,
                char_end=12,
                visibility=Visibility.PUBLIC,
            ),
            DocumentChunk(
                document_id=document.id,
                chunk_index=1,
                content="Draft chunk",
                char_start=13,
                char_end=24,
                visibility=Visibility.DRAFT,
            ),
        ]
    )
    await db_session.commit()

    public_chunks = await list_public_document_chunks(db_session)
    assert len(public_chunks) == 1
    assert public_chunks[0].content == "Public chunk"
